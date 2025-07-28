import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import io
from io import BytesIO
import os
import openai
from pathlib import Path
import base64
import streamlit.components.v1 as components

TAMANO_GRAFICO = (8, 4)
openai.api_key = os.getenv("OPENAI_API_KEY")

gris_o= "#7F7F7F"
gris_c= "#F2F2F2"
rojo="#D62828"
rosa_s= "#F7A9A8"
rosa_c="#FDE4E2"
palette=['#7F7F7F', '#F2F2F2', '#D62828', '#F7A9A8', '#FDE4E2']

st.markdown(
    """
    <style>
      /* Indicadores st.metric() */
      .element-container .stMetric {
        font-size: 12px !important;
        font-family: 'Calibri', 'Segoe UI', sans-serif !important;
        color: #7F7F7F !important;
      }
      .stMetric > div {
        color: #7F7F7F !important;
      }

      /* Botones normales (st.button) */
      .stButton > button {
        background-color: #FFFFFF !important;
        color: #D62828 !important;
        border: 2px solid #D62828 !important;
        font-family: 'Calibri', 'Segoe UI', sans-serif !important;
        font-size: 15px !important;
        font-weight: bold;
        border-radius: 8px;
        padding: 0.5rem 1.2rem;
        transition: all 0.3s ease;
      }

      .stButton > button:hover {
        background-color: #D62828 !important;
        color: #FFFFFF !important;
      }

      /* Botones de descarga (st.download_button) */
      div[data-testid="stDownloadButton"] > button {
        background-color: #FFFFFF !important;
        color: #D62828 !important;
        border: 2px solid #D62828 !important;
        font-family: 'Calibri', 'Segoe UI', sans-serif !important;
        font-size: 15px !important;
        font-weight: bold;
        border-radius: 8px;
        padding: 0.5rem 1.2rem;
        transition: all 0.3s ease;
      }

      div[data-testid="stDownloadButton"] > button:hover {
        background-color: #D62828 !important;
        color: #FFFFFF !important;
      }
    </style>
    """,
    unsafe_allow_html=True
)
def generar_analisis_gpt(prompt: str) -> str:
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "Eres un analista experto en seguros."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7,
        max_tokens=1200
    )
    return response.choices[0].message.content.strip()

def exportar_a_docx(texto: str) -> BytesIO:
    doc = Document()
    doc.add_heading("Informe Ejecutivo Generado por GPT", level=1)
    for linea in texto.split("\n"):
        doc.add_paragraph(linea)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
def encabezado_sin_icono(texto, nivel="h2"):
    estilo = {
        "h1": "font-size:28px; font-weight:bold;",
        "h2": "font-size:22px; font-weight:bold;",
        "h3": "font-size:18px;",
    }.get(nivel, "font-size:22px; font-weight:bold;")

    html = f"""
    <div style='margin-bottom:10px;'>
        <span style='{estilo} color:#D8272E; font-family:Calibri, sans-serif;'>{texto}</span>
    </div>
    """
    st.markdown(html, unsafe_allow_html=True)
    
def encabezado_con_icono(ruta_icono, texto, nivel="h2"):
    tamaños = {"h1": 28, "h2": 22, "h3": 18}
    estilo = {
        "h1": "font-size:28px; font-weight:bold;",
        "h2": "font-size:22px; font-weight:bold;",
        "h3": "font-size:18px;",
    }
    tamaño = tamaños.get(nivel, 22)
    estilo_texto = estilo.get(nivel, "font-size:22px;")

    icon_path = Path(ruta_icono)
    if not icon_path.exists():
        st.warning(f"⚠️ No se encontró el ícono en {ruta_icono}")
        return

    with open(icon_path, "rb") as f:
        icon_b64 = base64.b64encode(f.read()).decode()

    html = f"""
    <div style='display:flex; align-items:center; gap:10px; margin-bottom:10px;'>
        <img src="data:image/png;base64,{icon_b64}" style='height:{tamaño}px;' />
        <span style='{estilo_texto} color:#D8272E; font-family:Calibri, sans-serif;'>{texto}</span>
    </div>
    """
    st.markdown(html, unsafe_allow_html=True)

def render_tabla_html(df,height):
    table_html = """
    <style>
        .tabla-container {
            overflow-x: auto;
            padding: 0.5rem;
        }

        table.custom-table {
            border-collapse: separate;
            border-spacing: 0;
            width: 100%;
            font-family: 'Calibri', sans-serif;
            border: 1px solid #E0E0E0;
            border-radius: 10px;
            overflow: hidden;
        }

        table.custom-table th {
            background-color: #7F7F7F;
            color: white;
            font-weight: bold;
            text-align: center;
            padding: 10px;
            border-bottom: 1px solid #999999;
        }

        table.custom-table td {
            background-color: #F2F2F2;
            color: #7F7F7F;
            text-align: center;
            padding: 8px;
            border-bottom: 1px solid white;
        }

        table.custom-table tr.total-row td {
            background-color: #D62828;
            color: white;
            font-weight: bold;
            border-top: 1px solid white;
        }

    </style>
    <div class="tabla-container">
        <table class="custom-table">
            <thead>
                <tr>
    """

    for col in df.columns:
        table_html += f"<th>{col}</th>"
    table_html += "</tr></thead><tbody>"

    for index, row in df.iterrows():
        clase_fila = "total-row" if str(index).lower() == "total" else ""
        table_html += f'<tr class="{clase_fila}">'
        for val in row:
            table_html += f"<td>{val}</td>"
        table_html += "</tr>"

    table_html += "</tbody></table></div>"

    components.html(table_html, height=height, scrolling=True)
    
def mostrar_dashboard_analisis(pagados, pendientes, asegurados):
    
    if any(df is None or df.empty for df in [pagados, pendientes, asegurados]):
        st.warning("Por favor verifica que los DataFrames no estén vacíos.")
        return
    meses_orden = ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio',
                       'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre']
    mes_actual = datetime.now().month
    # Crear columnas necesarias en asegurados
    asegurados['FECHA'] = pd.to_datetime(asegurados['FECHA'], dayfirst=True, errors='coerce')
    asegurados['MES'] = asegurados['FECHA'].dt.month
    asegurados['AÑO'] = asegurados['FECHA'].dt.year
    
    # Crear columnas de comisiones por mes
    columnas_comision = [
        "COMISIÓN PRIMA VEHÍCULOS",
        "COMISIÓN CONCESIONARIO VEHÍCULOS",
        "COMISIÓN BROKER LIDERSEG VEHÍCULOS",
        "COMISIÓN BROKER INSURATLAN VEHÍCULOS"
    ]
    df_comisiones = asegurados.groupby(['AÑO', 'MES'])[columnas_comision].sum().reset_index()
    
    # Reclamos unificados
    mapeo_columnas = {
        'CIA. DE SEGUROS': 'COMPAÑÍA',
        'VALOR SINIESTRO': 'VALOR RECLAMO',
        'FECHA DE SINIESTRO': 'FECHA SINIESTRO'
    }
    pendientes_estandarizado = pendientes.rename(columns=mapeo_columnas)
    pagados['VALOR RECLAMO'] = pd.to_numeric(pagados['VALOR RECLAMO'], errors='coerce')
    pendientes_estandarizado['VALOR RECLAMO'] = pd.to_numeric(pendientes_estandarizado['VALOR RECLAMO'], errors='coerce')
    
    df_completo = pd.concat([
        pagados[['COMPAÑÍA', 'FECHA SINIESTRO', 'VALOR RECLAMO']],
        pendientes_estandarizado[['COMPAÑÍA', 'FECHA SINIESTRO', 'VALOR RECLAMO']]
    ], ignore_index=True).dropna(subset=['FECHA SINIESTRO'])
    
    df_completo['AÑO'] = pd.to_datetime(df_completo['FECHA SINIESTRO']).dt.year
    df_completo['MES'] = pd.to_datetime(df_completo['FECHA SINIESTRO']).dt.month
    
    valor_asegurado = asegurados.groupby(['ASEGURADORA', 'AÑO', 'MES']).agg(
        Prima_Vehiculos=('PRIMA VEHÍCULOS', 'sum')
    ).reset_index()
    
    reclamos_agrupados = df_completo.groupby(['COMPAÑÍA', 'AÑO', 'MES']).agg(
        Total_Reclamos=('VALOR RECLAMO', 'count'),
        Monto_Total_Reclamos=('VALOR RECLAMO', 'sum')
    ).reset_index().rename(columns={'COMPAÑÍA': 'ASEGURADORA'})
    
    df_siniestralidad = pd.merge(
        valor_asegurado,
        reclamos_agrupados,
        on=['ASEGURADORA', 'AÑO', 'MES'],
        how='left'
    ).fillna(0)
    
    df_siniestralidad['Siniestralidad'] = np.where(
        df_siniestralidad['Prima_Vehiculos'] > 0,
        df_siniestralidad['Monto_Total_Reclamos'] / df_siniestralidad['Prima_Vehiculos'],
        0
    )
    with st.container():
        # Borde visual mediante HTML
        st.markdown("<hr style='border:1px solid #7F7F7F; margin-top:0rem;'>", unsafe_allow_html=True)
        # fondo blanco con borde visual simulado como encabezado
        st.markdown(
            """
            <div style="border:1px solid #FFFFFF; border-radius:10px; padding:20px; background-color:#FFFFFF;">
            """,
            unsafe_allow_html=True
        )
        encabezado_con_icono("iconos/graficosubida.png", "Análisis de la Cuenta", "h1")
        seccion = st.radio(
            "Selecciona una sección:",
            ["Suma Asegurada", "Reclamos", "Siniestralidad","Comisiones por Canal","Generar Informe Ejecutivo"],
            horizontal=True
        )
        st.markdown("</div>", unsafe_allow_html=True)
    
        # Cierre visual del bloque
        st.markdown("<hr style='border:1px solid #7F7F7F; margin-top:1rem;'>", unsafe_allow_html=True)
            
    if seccion == "Suma Asegurada":
        # — Preparar columnas de fecha, mes y año —
        asegurados['FECHA'] = pd.to_datetime(
            asegurados['FECHA'], dayfirst=True, errors='coerce'
        )
        asegurados['MES'] = asegurados['FECHA'].dt.month
        asegurados['MES_NOMBRE'] = asegurados['MES'].apply(lambda x: meses_orden[x-1])
        asegurados['MES_NOMBRE'] = pd.Categorical(
            asegurados['MES_NOMBRE'], categories=meses_orden, ordered=True
        )
        asegurados['AÑO'] = asegurados['FECHA'].dt.year
    
        encabezado_con_icono("iconos/dinero.png", "Suma Asegurada", "h2")
    
        # — Sidebar de filtros —
        with st.sidebar:
            st.header("Configuración del Análisis de Suma Asegurada")
            aseguradoras = ['Todas'] + sorted(
                asegurados['ASEGURADORA'].dropna().unique().tolist()
            )
            aseguradora_sel = st.selectbox("Seleccionar Aseguradora", aseguradoras)
    
            años_disponibles = sorted(asegurados['AÑO'].dropna().unique())
            años_sel = st.multiselect(
                "Seleccionar Años", años_disponibles, default=años_disponibles
            )
    
            mostrar_graficos = st.multiselect(
                "Gráficos a mostrar",
                ["Distribuciones", "Evolución Anual", "Evolución Continua", "Tasa Mensual", "Top Marcas"]
            )
    
        # — Filtrar base según selección —
        df_base = asegurados.copy()
        if aseguradora_sel != 'Todas':
            df_base = df_base[df_base['ASEGURADORA'] == aseguradora_sel]
        if años_sel:
            df_base = df_base[df_base['AÑO'].isin(años_sel)]
        else:
            st.warning("Selecciona al menos un año")
            st.stop()
    
        # — Definir subconjuntos: Todos, Nuevos, Renovados —
        subconjuntos = {
            "Todos":    df_base,
            "Nuevos":   df_base[df_base['RENOVACION'] == 0],
            "Renovados":df_base[df_base['RENOVACION'] == 1],
        }
    
        # — Iterar sobre cada subconjunto y mostrar métricas, gráficos y tablas —
        for titulo, dfi in subconjuntos.items():
            encabezado_sin_icono(f"**{titulo}**", nivel="h2")
    
            # Métricas clave
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric(
                    "Suma Asegurada Total",
                    f"${dfi['VALOR ASEGURADO'].sum():,.2f}"
                )
            with col2:
                st.metric(
                    "Prima Total",
                    f"${dfi['PRIMA TOTAL VEHÍCULOS'].sum():,.2f}"
                )
            with col3:
                st.metric(
                    "Valor Promedio Asegurado",
                    f"${dfi['VALOR ASEGURADO'].mean():,.2f}"
                )
    
            # — Distribuciones —
            if "Distribuciones" in mostrar_graficos:
                encabezado_sin_icono("Distribuciones", "h3")
                c4, c5 = st.columns(2)
                with c4:
                    fig, ax = plt.subplots(figsize=TAMANO_GRAFICO)
                    sns.histplot(dfi['PRIMA TOTAL VEHÍCULOS'], kde=True, bins=30, ax=ax, color=rosa_s)
                    ax.set_title("Distribución de Prima Total")
                    ax.set_xlabel("Prima Total ($)")
                    ax.set_ylabel("Frecuencia")
                    st.pyplot(fig)
                with c5:
                    fig, ax = plt.subplots(figsize=TAMANO_GRAFICO)
                    sns.histplot(dfi['VALOR ASEGURADO'], kde=True, bins=30, ax=ax, color=rosa_s)
                    ax.set_title("Distribución de Valor Asegurado")
                    ax.set_xlabel("Valor Asegurado ($)")
                    ax.set_ylabel("Frecuencia")
                    st.pyplot(fig)
    
            # — Evolución Anual —
            if "Evolución Anual" in mostrar_graficos:
                encabezado_sin_icono("Evolución Anual", "h3")
                df_temp = dfi.pivot_table(
                    values='VALOR ASEGURADO',
                    index='MES',
                    columns='AÑO',
                    aggfunc='sum'
                ).reindex(range(1, 13))
                df_temp.index = pd.Categorical(
                    [meses_orden[m-1] for m in df_temp.index],
                    categories=meses_orden, ordered=True
                )
                for col in df_temp.columns:
                    mask = df_temp[col] != 0
                    if mask.any():
                        last = mask[mask].index[-1]
                        pos = df_temp.index.get_loc(last)
                        for i in range(pos+1, len(df_temp)):
                            if df_temp.iloc[i][col] == 0:
                                df_temp.iat[i, df_temp.columns.get_loc(col)] = np.nan
                fig, ax = plt.subplots(figsize=TAMANO_GRAFICO)
                df_temp.plot(ax=ax, marker='o')
                ax.set_title("Evolución Anual de la Suma Asegurada")
                ax.set_xlabel("Mes")
                ax.set_ylabel("Suma Asegurada")
                plt.xticks(rotation=45)
                st.pyplot(fig)
    
            # — Evolución Continua —
            if "Evolución Continua" in mostrar_graficos:
                encabezado_sin_icono("Evolución Continua", "h3")
                df_cont = dfi.copy()
                df_cont['Periodo'] = (
                    df_cont['MES'].apply(lambda x: meses_orden[x-1][:3])
                    + "-" + df_cont['AÑO'].astype(str)
                )
                orden = df_cont.sort_values(['AÑO', 'MES'])['Periodo'].unique()
                df_cont['Periodo'] = pd.Categorical(
                    df_cont['Periodo'], categories=orden, ordered=True
                )
                evo = df_cont.groupby('Periodo')['VALOR ASEGURADO'].sum().fillna(0)
                fig, ax = plt.subplots(figsize=TAMANO_GRAFICO)
                evo.plot(ax=ax, marker='o', legend=False, color=gris_o)
                ax.set_title("Evolución de la Suma Asegurada")
                ax.set_xlabel("Periodo")
                ax.set_ylabel("Suma Asegurada")
                plt.xticks(rotation=45)
                st.pyplot(fig)
    
            # — Tasa Mensual —
            if "Tasa Mensual" in mostrar_graficos:
                encabezado_sin_icono("Tasa Mensual", "h3")
                tasa = dfi.groupby(['AÑO', 'MES']).agg(
                    Prima_Total=('PRIMA TOTAL VEHÍCULOS', 'sum'),
                    Suma_Total=('VALOR ASEGURADO', 'sum')
                ).reset_index()
                tasa['Tasa'] = tasa['Prima_Total'] / tasa['Suma_Total'] * 100
                tasa['Periodo'] = (
                    tasa['MES'].apply(lambda x: meses_orden[x-1])
                    + "-" + tasa['AÑO'].astype(str)
                )
                tasa = tasa.sort_values(['AÑO', 'MES'])
                fig, ax = plt.subplots(figsize=TAMANO_GRAFICO)
                ax.plot(tasa['Periodo'], tasa['Tasa'], marker='o', color=gris_o)
                ax.set_title("Tasa Mensual de Prima vs. Suma Asegurada")
                ax.set_xlabel("Periodo")
                ax.set_ylabel("Tasa (%)")
                plt.xticks(rotation=45)
                st.pyplot(fig)
    
            # — Top Marcas —
            if "Top Marcas" in mostrar_graficos:
                encabezado_sin_icono("Top Marcas", "h3")
                top = dfi['MARCA'].value_counts().nlargest(10)
                fig, ax = plt.subplots(figsize=TAMANO_GRAFICO)
                top.plot(kind='barh', ax=ax, color=gris_o)
                ax.set_title("Top 10 Marcas Más Aseguradas")
                ax.set_xlabel("Cantidad")
                ax.set_ylabel("Marca")
                plt.xticks(rotation=45)
                st.pyplot(fig)
    
            # — Tablas Resumen —
            encabezado_sin_icono("Tablas Resumen", nivel="h3")
    
            # 1) Suma Asegurada por Aseguradora
            df_pri = (
                dfi.groupby(['ASEGURADORA', 'MES_NOMBRE', 'AÑO'])['VALOR ASEGURADO']
                .sum().reset_index()
            )
            pivot_pri = df_pri.pivot_table(
                index=['ASEGURADORA','MES_NOMBRE'],
                columns='AÑO',
                values='VALOR ASEGURADO',
                aggfunc='sum',
                margins=True,
                margins_name='Total'
            ).fillna(0).round(2).reset_index()
            pivot_pri['ASEGURADORA'] = pivot_pri['ASEGURADORA'].mask(
                pivot_pri['ASEGURADORA'].duplicated(), ''
            )
            render_tabla_html(pivot_pri, height=300)
    
            # 2) Unidades por Aseguradora
            df_uni = (
                dfi.groupby(['ASEGURADORA', 'AÑO'])['VALOR ASEGURADO']
                .count().reset_index(name='Unidades')
            )
            pivot_uni = df_uni.pivot_table(
                index='ASEGURADORA',
                columns='AÑO',
                values='Unidades',
                aggfunc='sum',
                margins=True,
                margins_name='Total'
            ).fillna(0).astype(int).reset_index()
            render_tabla_html(pivot_uni, height=200)
    
            # 3) Participación %
            part = dfi.pivot_table(
                index='ASEGURADORA',
                columns='AÑO',
                values='VALOR ASEGURADO',
                aggfunc='sum',
                margins=True,
                margins_name='Total'
            )
            part = (part.div(part.loc['Total'], axis=1) * 100).round(2).astype(str) + '%'
            render_tabla_html(part.reset_index(), height=200)
    
            # 4) Crecimiento Año a Año
            tot_ano = dfi.groupby('AÑO')['VALOR ASEGURADO'].sum().sort_index()
            df_crec = pd.DataFrame({
                'T. Valor Asegurado': tot_ano,
                'T. Valor LY': tot_ano.shift(1)
            })
            df_crec['%Crecimiento'] = (
                (df_crec['T. Valor Asegurado'] - df_crec['T. Valor LY'])
                / df_crec['T. Valor LY'] * 100
            ).round(2)
            df_crec = df_crec.dropna().reset_index()
            render_tabla_html(df_crec, height=200)
    
            # — Otras métricas por Mes y Año —
            df_extra = dfi.copy()
    
            # Unidades / Mes y Año
            df_unid = (
                df_extra.groupby(['MES_NOMBRE', 'AÑO'])
                .size().reset_index(name='Unidades')
            )
            pivot_unid = df_unid.pivot_table(
                index='MES_NOMBRE',
                columns='AÑO',
                values='Unidades',
                aggfunc='sum',
                margins=True,
                margins_name='Total'
            ).fillna(0).astype(int).reset_index()
            render_tabla_html(pivot_unid, height=300)
    
            # Promedio Valor Asegurado / Mes y Año
            df_val = (
                df_extra.groupby(['MES_NOMBRE', 'AÑO'])['VALOR ASEGURADO']
                .mean().reset_index()
            )
            pivot_val = df_val.pivot_table(
                index='MES_NOMBRE',
                columns='AÑO',
                values='VALOR ASEGURADO',
                aggfunc='mean',
                margins=True,
                margins_name='Total'
            ).round(2).fillna(0).reset_index()
            render_tabla_html(pivot_val, height=300)

    
    elif seccion == "Reclamos":
        encabezado_con_icono("iconos/reclamos.png", "Reclamos", "h2")
    
        # ——— Formateo de fechas y extracción de año base ———
        pagados['FECHA SINIESTRO']      = pd.to_datetime(pagados['FECHA SINIESTRO'],      errors='coerce')
        pendientes['FECHA DE SINIESTRO'] = pd.to_datetime(pendientes['FECHA DE SINIESTRO'], errors='coerce')
        pagados['BASE']     = pagados['FECHA SINIESTRO'].dt.year.astype('Int64')
        pendientes['BASE']  = pendientes['FECHA DE SINIESTRO'].dt.year.astype('Int64')
    
        # ——— Resumen para poblar multiselect de aseguradoras ———
        resumen_aseguradoras_total = (
            pagados
            .groupby('COMPAÑÍA')
            .agg(
                Media_Reclamo=('VALOR RECLAMO', 'mean'),
                Mediana_Reclamo=('VALOR RECLAMO', 'median'),
                Total_Reclamo=('VALOR RECLAMO', 'sum'),
                Media_Deducible=('DEDUCIBLE', 'mean')
            )
            .round(2)
        )
    
        # ——— Sidebar: año con opción "Todos" + selección de aseguradoras ———
        with st.sidebar:
            st.header("Configuración del Análisis de Reclamos")
    
            # Preparo lista de años dinámicamente y agrego "Todos"
            años_disponibles = sorted(pagados['BASE'].dropna().unique().tolist())
            opciones_años     = ["Todos"] + [str(a) for a in años_disponibles]
            año_analisis     = st.selectbox("Seleccionar Año", opciones_años, key="año_reclamos")
    
            # Multiselect de aseguradoras
            aseguradoras = resumen_aseguradoras_total.index.tolist()
            default_aseg = aseguradoras[:2] if len(aseguradoras) >= 2 else aseguradoras
            aseguradoras_seleccionadas = st.multiselect(
                "Selecciona Aseguradoras",
                options=aseguradoras,
                default=default_aseg,
                key="aseg_reclamos"
            )
    
        # ——— Filtrado según año (o todos) ———
        if año_analisis != "Todos":
            año_int = int(año_analisis)
            pagados_filtrados    = pagados[pagados['BASE'] == año_int].reset_index(drop=True)
            pendientes_filtrados = pendientes[pendientes['BASE'] == año_int].reset_index(drop=True)
        else:
            pagados_filtrados    = pagados.copy().reset_index(drop=True)
            pendientes_filtrados = pendientes.copy().reset_index(drop=True)
    
        # ——— Filtrado según aseguradora ———
        pagos_aseguradora_data      = pagados_filtrados[pagados_filtrados['COMPAÑÍA'].isin(aseguradoras_seleccionadas)]
        pendientes_aseguradora_data = pendientes_filtrados[pendientes_filtrados['CIA. DE SEGUROS'].isin(aseguradoras_seleccionadas)]
    
        # ——— Datos Generales ———
        encabezado_sin_icono("Datos Generales", "h2")
        encabezado_sin_icono("Reclamos Pagados", "h3")

        render_tabla_html(
            pagos_aseguradora_data[['COMPAÑÍA','VALOR RECLAMO','FECHA SINIESTRO','EVENTO']],
            height=250
        )
        encabezado_sin_icono("Reclamos Pendientes", "h3")

        render_tabla_html(
            pendientes_aseguradora_data[['CIA. DE SEGUROS','VALOR SINIESTRO','FECHA DE SINIESTRO','ESTADO ACTUAL']],
            height=250
        )
    
        encabezado_sin_icono("Distribución Temporal Continua", "h2")
        # 1) Crear columna PERIODO (timestamp al primer día de cada mes)
        pagos_aseguradora_data['PERIODO'] = (
            pagos_aseguradora_data['FECHA SINIESTRO']
            .dt.to_period('M')      # convertir a periodo mensual
            .dt.to_timestamp()      # volver a timestamp (primer día del mes)
        )
        
        # 2) Agrupar por PERIODO
        df_time = (
            pagos_aseguradora_data
            .groupby('PERIODO')
            .size()
            .reset_index(name='Reclamos')
            .sort_values('PERIODO')
        )
        # 3) Dibujar línea de serie temporal
        fig, ax = plt.subplots(figsize=TAMANO_GRAFICO)
        ax.plot(
                df_time['PERIODO'],
                df_time['Reclamos'],
                marker='o',
                color=gris_o,            # línea en gris_c
                markerfacecolor=rojo   # relleno del marcador también en gris_c
        )
        ax.set_title('Reclamos Mensuales (Continuos)')
        ax.set_ylabel('Número de Reclamos')
        ax.set_xlabel('Periodo')
        # Formatear etiquetas X como "Ene-2024", "Feb-2024", etc.
        xticks = df_time['PERIODO'][::3]
        ax.set_xticks(xticks)
        ax.set_xticklabels(xticks.dt.strftime('%b-%Y'), rotation=45, ha='right')        
        st.pyplot(fig)
    
        st.header("Análisis de Valores")
        encabezado_sin_icono("Análisis de Valores","h2")
        grafico_valores = st.radio("Elegir gráfico de análisis de valores", ["Histograma", "Boxplot", "Por Rangos"], horizontal=True)
        if grafico_valores == "Histograma":
            bins_hist = st.slider("Número de Bins", 10, 100, 30, 5)
            fig = plt.figure(figsize=TAMANO_GRAFICO)
            sns.histplot(pagos_aseguradora_data['VALOR RECLAMO'],color=gris_o ,bins=bins_hist, kde=True)
            st.pyplot(fig)
        elif grafico_valores == "Boxplot":
            fig = plt.figure(figsize=TAMANO_GRAFICO)
            sns.boxplot(x=pagos_aseguradora_data['VALOR RECLAMO'], color=rosa_s)
            st.pyplot(fig)
        elif grafico_valores == "Por Rangos":
            max_val = int(pagos_aseguradora_data['VALOR RECLAMO'].max())
            bin_size = st.slider("Tamaño del bin ($)", 2000, max_val, 2000, 500)
            bins = list(range(0, max_val + bin_size, bin_size))
            labels = [f"{bins[i]}-{bins[i+1]}" for i in range(len(bins)-1)]
            pagos_aseguradora_data['Rango'] = pd.cut(pagos_aseguradora_data['VALOR RECLAMO'], bins=bins, labels=labels, right=False)
            fig, ax = plt.subplots(figsize=TAMANO_GRAFICO)
            sns.countplot(y='Rango', data=pagos_aseguradora_data, order=labels, color=gris_o, ax=ax)
            st.pyplot(fig)
    
        # Nuevo selector de tipo de severidad
        encabezado_sin_icono("Análisis de Variables","h2")
        tipo_severidad = st.radio("Tipo de severidad", ["Promedio", "Total", "Frecuencia"], horizontal=True)

        def plot_severidad(tipo, campo, titulo):
            agrupado = pagos_aseguradora_data.groupby(campo)['VALOR RECLAMO']
            if tipo == "Promedio":
                datos = agrupado.mean()
                etiqueta = "Promedio ($)"
            elif tipo == "Total":
                datos = agrupado.sum()
                etiqueta = "Total ($)"
            else:  # Frecuencia
                datos = agrupado.count()
                etiqueta = "Cantidad de Reclamos"

            datos = datos.sort_values(ascending=False).head(10)
            fig, ax = plt.subplots(figsize=TAMANO_GRAFICO)
            sns.barplot(x=datos.values, y=datos.index, ax=ax, palette=palette)
            ax.set_title(titulo)
            ax.set_xlabel(etiqueta)
            st.pyplot(fig)

        cols = ["EVENTO", "TALLER DE REPARACION", "CIUDAD OCURRENCIA", "MARCA"]
        titulos = [
            "Severidad por Evento",
            "Severidad por Taller de Reparación",
            "Severidad por Ciudad",
            "Severidad por Marca"
        ]

        for col, titulo in zip(cols, titulos):
            plot_severidad(tipo_severidad, col, titulo)

        encabezado_sin_icono("Tablas Resumen", nivel="h2")

        # 1) Valor de Reclamos (US$) por Compañía, Mes y Año + Total
        encabezado_sin_icono("Valor de Reclamos (US$) por Compañía, Mes y Año + Total", nivel="h3")
        df_pri = pagos_aseguradora_data.copy()
        df_pri['MES_NOMBRE'] = df_pri['FECHA SINIESTRO'].dt.month.apply(lambda x: meses_orden[x-1])
        df_pri['AÑO']       = df_pri['FECHA SINIESTRO'].dt.year
        
        # Pivot con totales
        pivot_pri = (
            df_pri
            .groupby(['COMPAÑÍA','MES_NOMBRE','AÑO'])['VALOR RECLAMO']
            .sum()
            .reset_index()
        )
        pivot_pri['MES_NOMBRE'] = pd.Categorical(pivot_pri['MES_NOMBRE'], categories=meses_orden, ordered=True)
        pivot_pri = pivot_pri.pivot_table(
            index=['COMPAÑÍA','MES_NOMBRE'],
            columns='AÑO',
            values='VALOR RECLAMO',
            aggfunc='sum',
            margins=True,
            margins_name='Total'
        ).fillna(0).round(2)
        pivot_pri = pivot_pri.reset_index()
        pivot_pri['COMPAÑÍA'] = pivot_pri['COMPAÑÍA'].mask(pivot_pri['COMPAÑÍA'].duplicated(), '')
        render_tabla_html(pivot_pri, height=300)
        
        # 2) # de Reclamos por Compañía y Año
        encabezado_sin_icono("# de Reclamos por Compañía y Año", nivel="h3")
        df_uni = pagos_aseguradora_data.copy()
        df_uni['AÑO'] = df_uni['FECHA SINIESTRO'].dt.year
        
        pivot_unidades = (
            df_uni
            .groupby(['COMPAÑÍA','AÑO'])
            .size()
            .reset_index(name='Unidades')
            .pivot_table(
                index='COMPAÑÍA',
                columns='AÑO',
                values='Unidades',
                aggfunc='sum',
                margins=True,
                margins_name='Total'
            )
            .fillna(0)
            .astype(int)
            .reset_index()
        )
        render_tabla_html(pivot_unidades, height=200)
        
        # 3) Participación (%) por Compañía y Año
        encabezado_sin_icono("Participación de Valor Reclamos (%) por Compañía y Año", nivel="h3")
        pivot_part = pagos_aseguradora_data.copy()
        pivot_part['AÑO'] = pivot_part['FECHA SINIESTRO'].dt.year
        
        pivot_part = pivot_part.pivot_table(
            index='COMPAÑÍA',
            columns='AÑO',
            values='VALOR RECLAMO',
            aggfunc='sum',
            margins=True,
            margins_name='Total'
        )
        pivot_part = (pivot_part.div(pivot_part.loc['Total'], axis=1) * 100).round(2).astype(str) + '%'
        pivot_part = pivot_part.reset_index()
        render_tabla_html(pivot_part, height=200)
        
        # 4) Crecimiento Año a Año de Valor Reclamos
        encabezado_sin_icono("Crecimiento Año a Año de Valor Reclamos", nivel="h3")
        totales_ano = (
            pagos_aseguradora_data
            .assign(AÑO=pagos_aseguradora_data['FECHA SINIESTRO'].dt.year)
            .groupby('AÑO')['VALOR RECLAMO']
            .sum()
            .sort_index()
        )
        df_crec = pd.DataFrame({
            'T. Reclamos':    totales_ano,
            'T. Reclamos LY': totales_ano.shift(1)
        })
        df_crec['%Crecimiento'] = (
            (df_crec['T. Reclamos'] - df_crec['T. Reclamos LY'])
            / df_crec['T. Reclamos LY'] * 100
        ).round(2)
        df_crec = df_crec.dropna().reset_index().rename(columns={'index':'AÑO'})
        render_tabla_html(df_crec, height=200)
        
        # 5) # de Unidades / Mes y Año
        encabezado_sin_icono("# Unidades / Mes y Año", nivel="h3")
        df_extra = pagos_aseguradora_data.copy()
        
        # 1) Crear MES_NOMBRE y Año
        df_extra['MES_NOMBRE'] = df_extra['FECHA SINIESTRO'].dt.month.apply(lambda x: meses_orden[x-1])
        df_extra['AÑO']        = df_extra['FECHA SINIESTRO'].dt.year
        
        # 2) Categorizar MES_NOMBRE con orden
        df_extra['MES_NOMBRE'] = pd.Categorical(
            df_extra['MES_NOMBRE'],
            categories=meses_orden,
            ordered=True
        )
        
        # 3) Pivot para unidades
        pivot_unid = (
            df_extra
            .groupby(['MES_NOMBRE','AÑO'])
            .size()
            .reset_index(name='Unidades')
            .pivot_table(
                index='MES_NOMBRE',
                columns='AÑO',
                values='Unidades',
                aggfunc='sum',
                margins=True,
                margins_name='Total'
            )
            .fillna(0)
            .astype(int)
        )
        
        # 4) Reindexar para forzar orden cronológico + Total al final
        pivot_unid = pivot_unid.reindex(index=meses_orden + ['Total'])
        
        # 5) Volver a columna y mostrar
        pivot_unid = pivot_unid.reset_index()
        render_tabla_html(pivot_unid, height=300)
        
        
        # ——— Promedio Valor Reclamo (US$) / Mes y Año ———
        encabezado_sin_icono("Promedio Valor Reclamo (US$) / Mes y Año", nivel="h3")
        # (df_extra ya tiene MES_NOMBRE categórico)
        pivot_prom = (
            df_extra
            .groupby(['MES_NOMBRE','AÑO'])['VALOR RECLAMO']
            .mean()
            .round(2)
            .reset_index()
            .pivot_table(
                index='MES_NOMBRE',
                columns='AÑO',
                values='VALOR RECLAMO',
                aggfunc='mean',
                margins=True,
                margins_name='Total'
            )
            .fillna(0)
        )
        
        # Reindexar igual que arriba
        pivot_prom = pivot_prom.reindex(index=meses_orden + ['Total'])
        pivot_prom = pivot_prom.reset_index()
        render_tabla_html(pivot_prom, height=300)
        encabezado_sin_icono("Métricas Clave", "h2")

        # ——— KPIs ———
        # Cálculos rápidos
        total_pagados    = len(pagos_aseguradora_data)
        total_pendientes = len(pendientes_aseguradora_data)
        valor_total      = pagos_aseguradora_data['VALOR RECLAMO'].sum()
        valor_promedio   = pagos_aseguradora_data['VALOR RECLAMO'].mean()
        valor_max        = pagos_aseguradora_data['VALOR RECLAMO'].max()
        
        # Mostrar como métricas en tres columnas
        col1, col2, col3 = st.columns(3)
        col1.metric("Reclamos Pagados",    f"{total_pagados}",    f"{total_pendientes} pendientes")
        col2.metric("Valor Total Pagado",   f"${valor_total:,.2f}", f"Promedio ${valor_promedio:,.2f}")
        col3.metric("Mayor Reclamo",        f"${valor_max:,.2f}")
    
        encabezado_sin_icono("Generar Informe Anual","h2")

        if st.button("Generar Informe"):
                # ——— Crear columna PERIODO en pagados_filtrados ———
            pagados_filtrados['PERIODO'] = (
                pagados_filtrados['FECHA SINIESTRO']
                .dt.to_period('M')      # convierte a periodo mensual “2024-01”, etc.
                .dt.to_timestamp()      # vuelve a timestamp (1er día del mes)
            )
        
            # ——— Pivot: ahora por PERIODO en lugar de MES ———
            resumen_periodo = pagados_filtrados.pivot_table(
                values='VALOR RECLAMO',
                index='PERIODO',
                columns='COMPAÑÍA',
                aggfunc=['sum', 'count'],
                fill_value=0,
                margins=True,
                margins_name='Total'
            )
            # Aplanar multiíndice de columnas
            resumen_periodo.columns = [f"{aggfunc} {aseg}" for aggfunc, aseg in resumen_periodo.columns]
            talleres = pagados_filtrados.pivot_table(values='VALOR RECLAMO', index='TALLER DE REPARACION', columns='COMPAÑÍA', aggfunc='count', fill_value=0)
            causas = pagados_filtrados.pivot_table(values='VALOR RECLAMO', index='EVENTO', aggfunc=['sum', 'count'], fill_value=0)
            causas.columns = ['Total_Reclamo', 'Cantidad_Reclamos']
            pendientes_estado = pendientes_filtrados.pivot_table(values='VALOR SINIESTRO', index='ESTADO ACTUAL', columns='CIA. DE SEGUROS', aggfunc='count', fill_value=0)

            # Severidad resumen
            def resumen_severidad(df, campo):
                return pd.DataFrame({
                    'Severidad Promedio': df.groupby(campo)['VALOR RECLAMO'].mean().round(2),
                    'Severidad Total': df.groupby(campo)['VALOR RECLAMO'].sum().round(2),
                    'Frecuencia': df.groupby(campo)['VALOR RECLAMO'].count()
                }).sort_values('Severidad Total', ascending=False)

            sev_evento = resumen_severidad(pagos_aseguradora_data, 'EVENTO')
            sev_taller = resumen_severidad(pagos_aseguradora_data, 'TALLER DE REPARACION')
            sev_ciudad = resumen_severidad(pagos_aseguradora_data, 'CIUDAD OCURRENCIA')
            sev_marca = resumen_severidad(pagos_aseguradora_data, 'MARCA')

            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                resumen_periodo.to_excel(writer, sheet_name='Resumen Por Periodo')
                talleres.to_excel(writer, sheet_name='Talleres')
                causas.to_excel(writer, sheet_name='Causas')
                pendientes_estado.to_excel(writer, sheet_name='Pendientes')
                sev_evento.to_excel(writer, sheet_name='Severidad Eventos')
                sev_taller.to_excel(writer, sheet_name='Severidad Talleres')
                sev_ciudad.to_excel(writer, sheet_name='Severidad Ciudades')
                sev_marca.to_excel(writer, sheet_name='Severidad Marcas')
            output.seek(0)

            st.download_button(
                label="Descargar Reporte",
                data=output,
                file_name=f"Reporte_Reclamos_{año_analisis}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
    
    elif seccion == "Siniestralidad":
        encabezado_con_icono("iconos/graficosubida.png","Siniestralidad Mensual por Aseguradora","h2")
    
        # Estándar de nombres desde pendientes
        mapeo_columnas = {
            'CIA. DE SEGUROS': 'COMPAÑÍA',
            'VALOR SINIESTRO': 'VALOR RECLAMO',
            'FECHA DE SINIESTRO': 'FECHA SINIESTRO'
        }
        pendientes_estandarizado = pendientes.rename(columns=mapeo_columnas)
    
        # Asegurar tipos correctos
        pagados['FECHA SINIESTRO'] = pd.to_datetime(pagados['FECHA SINIESTRO'], errors='coerce')
        pagados['VALOR RECLAMO'] = pd.to_numeric(pagados['VALOR RECLAMO'], errors='coerce')
    
        pendientes_estandarizado['FECHA SINIESTRO'] = pd.to_datetime(pendientes_estandarizado['FECHA SINIESTRO'], errors='coerce')
        pendientes_estandarizado['VALOR RECLAMO'] = pd.to_numeric(pendientes_estandarizado['VALOR RECLAMO'], errors='coerce')
    
        # Unir reclamos pagados + pendientes
        df_completo = pd.concat([
            pagados[['COMPAÑÍA', 'FECHA SINIESTRO', 'VALOR RECLAMO']],
            pendientes_estandarizado[['COMPAÑÍA', 'FECHA SINIESTRO', 'VALOR RECLAMO']]
        ], ignore_index=True).dropna(subset=['FECHA SINIESTRO'])
    
        # Generar columnas de año y mes
        df_completo['FECHA_SINIESTRO'] = pd.to_datetime(df_completo['FECHA SINIESTRO'], errors='coerce')
        df_completo['AÑO'] = df_completo['FECHA_SINIESTRO'].dt.year.astype('Int64')
        df_completo['MES'] = df_completo['FECHA_SINIESTRO'].dt.month.astype('Int64')
    
        # Valor asegurado por mes
        valor_asegurado = asegurados.groupby(['ASEGURADORA', 'AÑO', 'MES']).agg(
            Prima_Vehiculos=('PRIMA VEHÍCULOS', 'sum')
        ).reset_index()
    
        # Reclamos agregados
        reclamos_agrupados = df_completo.groupby(['COMPAÑÍA', 'AÑO', 'MES']).agg(
            Total_Reclamos=('VALOR RECLAMO', 'count'),
            Monto_Total_Reclamos=('VALOR RECLAMO', 'sum')
        ).reset_index().rename(columns={'COMPAÑÍA': 'ASEGURADORA'})
    
        # Unión con asegurados
        df_siniestralidad = pd.merge(
            valor_asegurado,
            reclamos_agrupados,
            on=['ASEGURADORA', 'AÑO', 'MES'],
            how='left'
        ).fillna(0)
    
        df_siniestralidad['Siniestralidad'] = np.where(
            df_siniestralidad['Prima_Vehiculos'] > 0,
            df_siniestralidad['Monto_Total_Reclamos'] / df_siniestralidad['Prima_Vehiculos'],
            0
        )
    
        df_siniestralidad['PERIODO'] = df_siniestralidad['AÑO'].astype(str) + '-' + df_siniestralidad['MES'].astype(str).str.zfill(2)
        df_siniestralidad['FECHA'] = pd.to_datetime(df_siniestralidad['PERIODO'] + '-01', errors='coerce')
    
        # Filtros
        col1, col2 = st.columns(2)
        with col1:
            años_disponibles = sorted(df_siniestralidad['AÑO'].dropna().unique(), reverse=True)
            año_sel = st.selectbox("Año de análisis", options=['Todos'] + list(años_disponibles), key='sini_año')
        with col2:
            aseguradora_sel = st.selectbox("Aseguradora", options=['Todas'] + sorted(df_siniestralidad['ASEGURADORA'].unique()), key='sini_aseguradora')
    
        # Filtro por año
        if año_sel != 'Todos':
            df_filtrado = df_siniestralidad[df_siniestralidad['AÑO'] == int(año_sel)]
        else:
            df_filtrado = df_siniestralidad.copy()
    
        # Filtro por aseguradora
        if aseguradora_sel != 'Todas':
            df_filtrado = df_filtrado[df_filtrado['ASEGURADORA'] == aseguradora_sel]
        else:
            group_cols = ['PERIODO', 'AÑO', 'MES', 'FECHA'] if año_sel == 'Todos' else ['PERIODO', 'AÑO', 'MES']
            df_filtrado = df_filtrado.groupby(group_cols).agg({
                'Prima_Vehiculos': 'sum',
                'Total_Reclamos': 'sum',
                'Monto_Total_Reclamos': 'sum'
            }).reset_index()
            df_filtrado['Siniestralidad'] = np.where(
                df_filtrado['Prima_Vehiculos'] > 0,
                df_filtrado['Monto_Total_Reclamos'] / df_filtrado['Prima_Vehiculos'],
                0
            )
            df_filtrado['FECHA'] = pd.to_datetime(df_filtrado['PERIODO'] + '-01', errors='coerce')
    
        # Gráfico
        df_filtrado = df_filtrado.sort_values('FECHA')
        fig, ax = plt.subplots(figsize=(14, 6))
        sns.lineplot(data=df_filtrado, x='PERIODO', y='Siniestralidad', marker='o', color=rojo, ax=ax)
        ax.set_ylabel("Ratio Siniestralidad", color='#E53935')
        ax.tick_params(axis='y', colors='#E53935')
    
        num_periodos = len(df_filtrado['PERIODO'].unique())
        rotation = 45 if num_periodos > 6 else 0
        step = max(1, num_periodos // 12)
        ax.set_xticks(df_filtrado['PERIODO'].unique()[::step])
        ax.set_xticklabels(df_filtrado['PERIODO'].unique()[::step], rotation=rotation)
    
        if aseguradora_sel != 'Todas':
            ax2 = ax.twinx()
            sns.barplot(data=df_filtrado, x='PERIODO', y='Monto_Total_Reclamos', color=rosa_s, alpha=0.3, ax=ax2)
            ax2.set_ylabel("Monto Total Reclamos ($)", color=rosa_s)
            ax2.tick_params(axis='y', colors=rosa_s)
    
        titulo = f"Siniestralidad {'por Aseguradora' if aseguradora_sel != 'Todas' else 'Acumulada'}"
        titulo += f" ({'Histórico Completo' if año_sel == 'Todos' else año_sel})"
        plt.title(titulo)
        st.pyplot(fig)
    

        encabezado_sin_icono("Datos Detallados",nivel="h2")
        columnas = ['PERIODO', 'ASEGURADORA', 'Prima_Vehiculos', 'Total_Reclamos', 'Monto_Total_Reclamos', 'Siniestralidad'] if aseguradora_sel != 'Todas' else ['PERIODO', 'Prima_Vehiculos', 'Total_Reclamos', 'Monto_Total_Reclamos', 'Siniestralidad']
        df_tabla = df_filtrado[columnas].copy()
        df_tabla['Prima_Vehiculos'] = df_tabla['Prima_Vehiculos'].map("${:,.2f}".format)
        df_tabla['Monto_Total_Reclamos'] = df_tabla['Monto_Total_Reclamos'].map("${:,.2f}".format)
        df_tabla['Total_Reclamos'] = df_tabla['Total_Reclamos'].map("{:,.0f}".format)
        df_tabla['Siniestralidad'] = df_tabla['Siniestralidad'].map("{:.2%}".format)
        
        render_tabla_html(df_tabla,height=450)
        encabezado_sin_icono("Indicadores Clave",nivel="h2")

        if not df_filtrado.empty:
            ultimo_mes = df_filtrado.iloc[-1]
            primer_mes = df_filtrado.iloc[0]
        
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Periodo Inicial", primer_mes['PERIODO'])
            with col2:
                st.metric("Periodo Final", ultimo_mes['PERIODO'])
            with col3:
                st.metric("Total Periodos", len(df_filtrado))
        
            col4, col5, col6 = st.columns(3)
            with col4:
                st.metric("Prima Vehículos Promedio", f"${df_filtrado['Prima_Vehiculos'].mean():,.2f}")
            with col5:
                st.metric("Siniestralidad Promedio", f"{df_filtrado['Siniestralidad'].mean():.2%}")
            with col6:
                st.metric("Total Reclamos", f"{df_filtrado['Total_Reclamos'].sum():,.0f}")
            
    elif seccion == "Comisiones por Canal":
        encabezado_con_icono("iconos/dinero.png","Análisis de Comisiones por Canal","h2")
    
        # Asegurar formato de fecha y crear columnas de año y mes
        asegurados['FECHA'] = pd.to_datetime(asegurados['FECHA'], dayfirst=True, errors='coerce')
        asegurados['MES'] = asegurados['FECHA'].dt.month
        asegurados['AÑO'] = asegurados['FECHA'].dt.year
    
        with st.sidebar:
            st.header("Filtros de Comisiones")
            
            aseguradoras = ['Todas'] + sorted(asegurados['ASEGURADORA'].dropna().unique())
            aseguradora_sel = st.selectbox("Seleccionar Aseguradora", aseguradoras, key="aseg_comisiones")
    
            años_disponibles = sorted(asegurados['AÑO'].dropna().unique())
            años_sel = st.multiselect("Seleccionar Años", años_disponibles, default=años_disponibles, key="años_comisiones")
    
        df_com = asegurados.copy()
        if aseguradora_sel != 'Todas':
            df_com = df_com[df_com['ASEGURADORA'] == aseguradora_sel]
    
        if años_sel:
            df_com = df_com[df_com['AÑO'].isin(años_sel)]
        else:
            st.warning("Selecciona al menos un año")
            st.stop()
    
        columnas_comision = [
            "COMISIÓN PRIMA VEHÍCULOS",
            "COMISIÓN CONCESIONARIO VEHÍCULOS",
            "COMISIÓN BROKER LIDERSEG VEHÍCULOS",
            "COMISIÓN BROKER INSURATLAN VEHÍCULOS"
        ]
    
        df_comisiones = df_com.groupby(['AÑO', 'MES'])[columnas_comision].sum().reset_index()
    
        if df_comisiones.empty:
            st.warning("No hay datos de comisiones para los filtros seleccionados.")
            return
    
        df_comisiones['Periodo'] = df_comisiones['MES'].apply(lambda x: meses_orden[x - 1]) + '-' + df_comisiones['AÑO'].astype(str)
        df_comisiones = df_comisiones.sort_values(['AÑO', 'MES'])
        df_comisiones['Periodo'] = pd.Categorical(df_comisiones['Periodo'], categories=df_comisiones['Periodo'].unique(), ordered=True)
        df_comisiones.set_index('Periodo', inplace=True)
    
        # Total comisiones
        if not df_comisiones.empty:
            encabezado_sin_icono("Total Comisiones Pagadas",nivel="h2")
            total_comisiones = df_comisiones[columnas_comision].sum().sum()
            st.metric("Total USD", f"${total_comisiones:,.2f}")
    
            # Gráfico apilado por canal
            encabezado_sin_icono("Evolución de Comisiones por Canal",nivel="h2")
            fig, ax = plt.subplots(figsize=TAMANO_GRAFICO)
            paleta_colores = ['#7F7F7F', '#C5C5C5', '#D62828', '#F7A9A8']  # Usa solo 4 si tienes 4 canales

            df_comisiones[columnas_comision].plot(kind='bar',color=paleta_colores ,stacked=True, ax=ax)
            ax.set_ylabel("USD ($)")
            ax.set_title("Pago de Comisiones por Canal y Mes")
            ax.legend(title="Canal", bbox_to_anchor=(1.05, 1), loc='upper left')
            plt.xticks(rotation=45)
            st.pyplot(fig)
    
            # Gráfico individual por canal seleccionado
            encabezado_sin_icono("Comisiones por Canal - Individual",nivel="h2")
            canal_seleccionado = st.selectbox("Selecciona el canal a visualizar:", columnas_comision)
            if canal_seleccionado in df_comisiones.columns:
                fig, ax = plt.subplots(figsize=TAMANO_GRAFICO)
                ax.plot(df_comisiones.index, df_comisiones[canal_seleccionado], color=gris_o ,marker='o', label=canal_seleccionado)
                ax.set_title(f"Evolución de {canal_seleccionado}")
                ax.set_xlabel("Periodo")
                ax.set_ylabel("USD ($)")
                plt.xticks(rotation=45)
                st.pyplot(fig)
            else:
                st.warning(f"El canal '{canal_seleccionado}' no está disponible en los datos.")
    
            # Tabla detallada
            encabezado_sin_icono("Tabla de Comisiones Detallada", nivel="h3")
            # 1) Resetear índice para que 'Periodo' vuelva a ser columna
            df_tabla = (
                df_comisiones
                .reset_index()[['Periodo'] + columnas_comision]  # traemos Periodo + comisiones
                .copy()
                .round(2)
            )
            render_tabla_html(df_tabla, height=250)
            # Exportar a Excel
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                df_comisiones.reset_index()[['Periodo'] + columnas_comision].to_excel(writer, sheet_name='Comisiones Mensuales', index=False)
            output.seek(0)
            st.download_button(
                label="Descargar Comisiones en Excel",
                data=output,
                file_name="comisiones_por_canal.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
    elif seccion == "Generar Informe Ejecutivo":
        encabezado_con_icono("iconos/informe.png","Análisis Final Automatizado con AI",nivel="h2")
        datos_ok = all(
            df is not None and isinstance(df, pd.DataFrame) and not df.empty
            for df in [pagados, pendientes, asegurados, df_comisiones, df_siniestralidad]
        )
    
        if not datos_ok:
            st.warning("No se han podido cargar correctamente todos los datos requeridos para el análisis.")
        else:
            if st.button("Generar Análisis con AI"):
                with st.spinner("Consultando modelo de AI..."):
                    prompt = f"""
    Analiza los siguientes datos brevemente:
    
    1. **Comisiones Totales por Mes**:
    {df_comisiones[columnas_comision].tail(3).to_string(index=True)}
    
    2. **Reclamos Pagados por Aseguradora (muestra):**
    {pagados[['COMPAÑÍA','VALOR RECLAMO','EVENTO','TALLER DE REPARACION','CONCESIONARIO SISTEMA']].head(3).to_string(index=False)}
    
    3. **Siniestralidad por Aseguradora**:
    {df_siniestralidad[['ASEGURADORA','AÑO','MES','Siniestralidad']].dropna().tail(3).to_string(index=False)}
    
    4. **Suma Asegurada por Marca (muestra):**
    {asegurados[['MARCA','VALOR ASEGURADO','FECHA']].dropna().tail(3).to_string(index=False)}
    
    Escribe un **informe ejecutivo en español**, que incluya (excluir el mes en vigencia ya que este todavia no cuenta con informacion completa):
    - Resumen general.
    - Análisis de comisiones: meses altos/bajos.
    - Aseguradoras con más reclamos y eventos más frecuentes. 
    - Comportamiento de concesionarios y talleres. (Ponle mucho enfasis a este punto)
    - Siniestralidad: ¿alta/baja?, ¿mejores y peores aseguradoras?
    - Marcas más aseguradas y evolución de la suma asegurada.
    - Cierre con recomendaciones o insights accionables.
    """
                    try:
                        analisis = generar_analisis_gpt(prompt)
                        st.markdown(analisis)
                        docx_file = exportar_a_docx(analisis)
                        st.download_button(
                            label="Descargar Informe en Word",
                            data=docx_file,
                            file_name="informe_gpt.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except Exception as e:
                        st.error(f"Error al generar análisis: {e}")





